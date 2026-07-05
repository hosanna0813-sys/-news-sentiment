"""測試：Word 匯出（規格十四）"""
from __future__ import annotations

import os
from app.models.topic import Topic, Stance
from app.models.news import NewsItem
from app.models.settings import WordExportSettings
from app.exporters.word_exporter import export_daily_report


def test_export_creates_valid_docx_with_expected_sections(tmp_path):
    topic = Topic(
        topic_id="t1", topic_name="測試議題：某政策爭議",
        summary_150="150字摘要內容。", summary_300="300字摘要內容。",
        development_progress="事件持續發展中。", core_disputes="核心爭點測試。",
        key_actors="主要行動者說明。", possible_impact="可能後續影響說明。",
        has_identifiable_stance=True,
    )
    news = [
        NewsItem(row_id="r1", title="測試新聞一", source="中央社", published_at="2026-07-01",
                 url="http://example.com/1", body_text="正文..."),
        NewsItem(row_id="r2", title="測試新聞二", source="聯合報", published_at="2026-07-02",
                 url="http://example.com/2", body_text="正文2..."),
    ]
    stance = Stance(stance_id="s1", topic_id="t1", stance_type="支持", speaker="某官員",
                     organization="某部會", claim="支持該政策", evidence_news_id="r1",
                     evidence_excerpt="官員表示支持")
    missing_body_news = [NewsItem(row_id="r3", title="未取得正文新聞", source="蘋果日報",
                                   body_fetch_status="失敗", body_fetch_detail="404")]

    out_path = tmp_path / "report.docx"
    settings = WordExportSettings()
    result_path = export_daily_report(
        str(out_path), [topic], {"t1": news}, {"t1": [stance]}, missing_body_news, settings,
    )

    assert os.path.exists(result_path)
    assert os.path.getsize(result_path) > 0

    from docx import Document
    doc = Document(result_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)

    assert "測試議題：某政策爭議" in full_text
    assert "150字摘要內容" in full_text
    assert "測試新聞一" in full_text
    assert "測試新聞二" in full_text
    assert "未取得可用正文之新聞清單" in full_text
    assert "未取得正文新聞" in full_text


def test_stance_section_hidden_when_no_identifiable_stance(tmp_path):
    """規格十二：若議題群無明確立場，Word 報告不應顯示立場區塊"""
    topic = Topic(topic_id="t2", topic_name="純事實議題", summary_150="摘要",
                   has_identifiable_stance=False)
    news = [NewsItem(row_id="r1", title="純事實新聞", source="中央社", body_text="正文")]

    out_path = tmp_path / "report2.docx"
    export_daily_report(str(out_path), [topic], {"t2": news}, {"t2": []}, [], WordExportSettings())

    from docx import Document
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "主要論述與立場" not in full_text


def test_simple_topic_list_export(tmp_path):
    """簡易清單：議題分組＋標題＋連結，不含摘要；無 url 新聞不產生連結"""
    from app.exporters.word_exporter import export_simple_topic_list

    t1 = Topic(topic_id="t1", topic_name="議題一：政策爭議", summary_150="不該出現的摘要內容。")
    t_empty = Topic(topic_id="t0", topic_name="空議題不應出現")  # 排中間，驗證不佔編號
    t2 = Topic(topic_id="t2", topic_name="議題二：治安事件")
    news = {
        "t1": [
            NewsItem(row_id="r1", title="新聞甲", source="中央社", url="http://example.com/a"),
            NewsItem(row_id="r2", title="新聞乙（無來源無連結）", source="", url=""),
        ],
        "t2": [NewsItem(row_id="r3", title="新聞丙", source="自由", url="http://example.com/c")],
        "t0": [],
    }

    out = tmp_path / "simple.docx"
    result = export_simple_topic_list(str(out), [t1, t_empty, t2], news, WordExportSettings())
    assert os.path.exists(result)

    from docx import Document
    doc = Document(result)
    texts = [p.text for p in doc.paragraphs]
    full_text = "\n".join(texts)
    assert "1. 議題一：政策爭議" in texts        # 議題有編號
    assert "2. 議題二：治安事件" in texts        # 空議題不佔編號
    assert "中央社-新聞甲" in texts              # 來源-標題 前綴
    assert "自由-新聞丙" in texts
    assert "新聞乙（無來源無連結）" in texts     # 無來源時只印標題
    assert not any(t.startswith("-") for t in texts)  # 不出現孤兒的 -
    assert "不該出現的摘要內容" not in full_text   # 不含摘要
    assert "空議題不應出現" not in full_text       # 空議題跳過
    xml = doc.element.xml
    assert xml.count("</w:hyperlink>") == 2        # 只有兩則有連結
    assert "http://example.com/a" in xml
