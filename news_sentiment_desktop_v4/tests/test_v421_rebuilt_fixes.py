"""測試：V4.2.1 依 HANDOFF.md 重建的六項修改

1. BatchJobWorker cleanup_fn（worker 執行緒 finally 收尾，Playwright EPIPE 修正）
2. Prompt 編輯器佔位符工具（extract_placeholders）
3. Word 匯出版型（標楷體 eastAsia／摺疊／超連結／key_actors／無標籤前綴）
4. 模型輸出清洗（strip_model_artifacts / strip_artifacts_deep + gateway 套用）
5. 摘要 180 字限制（truncate_at_sentence + prompt 要求）
6. 議題調整頁正文編輯（human_edit_body 資料流）
"""
from __future__ import annotations

import threading
import time

from app.utils.text_utils import (
    strip_model_artifacts, strip_artifacts_deep, truncate_at_sentence, extract_placeholders,
)


# ---------------------------------------------------------------------------
# 修改 4：模型輸出清洗
# ---------------------------------------------------------------------------
def test_strip_model_artifacts_removes_leaked_tags():
    text = '事件持續發展。</summary_150>後續<parameter name="key_actors">仍待觀察。<tool_call/>'
    cleaned = strip_model_artifacts(text)
    assert "</summary_150>" not in cleaned
    assert "<parameter" not in cleaned
    assert "<tool_call/>" not in cleaned
    assert "事件持續發展。" in cleaned
    assert "仍待觀察。" in cleaned


def test_strip_model_artifacts_converts_literal_backslash_n():
    text = "第一位：甲\\n第二位：乙"
    cleaned = strip_model_artifacts(text)
    assert "\\n" not in cleaned
    assert "第一位：甲\n第二位：乙" == cleaned


def test_strip_model_artifacts_keeps_normal_angle_brackets():
    # 一般比較符號與非屬性語法的角括號內容不應被誤刪
    text = "支出 3<5 億元，民眾稱「<不合理>」"
    cleaned = strip_model_artifacts(text)
    assert "3<5" in cleaned


def test_strip_artifacts_deep_cleans_nested_structures():
    data = {
        "summary_150": "摘要。</summary_150>",
        "stances": [{"claim": "主張\\n第二行"}],
        "count": 3,
    }
    cleaned = strip_artifacts_deep(data)
    assert cleaned["summary_150"] == "摘要。"
    assert cleaned["stances"][0]["claim"] == "主張\n第二行"
    assert cleaned["count"] == 3


def test_gateway_tool_use_output_is_cleaned(fake_anthropic_module):
    """gateway 的 tool use 輸出路徑套用清洗"""
    from app.services.ai.model_gateway import ModelGateway

    class FakeUsage:
        input_tokens = 10
        output_tokens = 20

    class FakeToolBlock:
        type = "tool_use"
        input = {"summary_150": "內容。</summary_150>", "note": "第一行\\n第二行"}

    class FakeResp:
        content = [FakeToolBlock()]
        stop_reason = "tool_use"
        usage = FakeUsage()

    gw = ModelGateway(api_key_provider=lambda: "sk-test",
                       task_model_lookup=lambda t: {"model_id": "m", "max_tokens": 100})
    gw._create_message = lambda *a, **k: FakeResp()
    result = gw.call_with_tool("topic_summarization", "sys", "user", "tool",
                                {"type": "object", "properties": {}})
    assert result.data["summary_150"] == "內容。"
    assert result.data["note"] == "第一行\n第二行"


# ---------------------------------------------------------------------------
# 修改 5：摘要 180 字限制
# ---------------------------------------------------------------------------
def test_truncate_at_sentence_short_text_unchanged():
    assert truncate_at_sentence("短摘要。", 180) == "短摘要。"


def test_truncate_at_sentence_cuts_at_sentence_ending():
    text = "第一句完整敘述。" * 30  # 240 字
    result = truncate_at_sentence(text, 180)
    assert len(result) <= 180
    assert result.endswith("。")


def test_truncate_at_sentence_hard_cut_when_no_punctuation():
    text = "無標點" * 100
    result = truncate_at_sentence(text, 180)
    assert len(result) == 180


def test_summarization_prompt_contains_180_limit_and_actor_line_rule():
    from app.prompts.summarization_prompt import (
        SUMMARIZATION_SYSTEM_PROMPT, SUMMARIZATION_USER_TEMPLATE,
    )
    assert "180" in SUMMARIZATION_SYSTEM_PROMPT
    assert "不可留空" in SUMMARIZATION_SYSTEM_PROMPT
    assert "每位" in SUMMARIZATION_SYSTEM_PROMPT      # key_actors 每位一行
    assert "summary_150" in SUMMARIZATION_USER_TEMPLATE  # 欄位名稱明確對應


# ---------------------------------------------------------------------------
# 修改 2：Prompt 編輯器佔位符工具
# ---------------------------------------------------------------------------
def test_extract_placeholders():
    assert extract_placeholders("議題：{topic_name}\n{topic_news_json}") == {
        "topic_name", "topic_news_json"}
    assert extract_placeholders("沒有佔位符") == set()
    assert extract_placeholders("") == set()


# ---------------------------------------------------------------------------
# 修改 1：BatchJobWorker cleanup_fn
# ---------------------------------------------------------------------------
def _make_worker(job_repo, batch_repo, process_fn, cleanup_fn, batches=None):
    from app.workers.batch_job_worker import BatchJobWorker
    return BatchJobWorker(
        job_type="scraping", item_batches=batches or [["a", "b"], ["c"]],
        process_batch_fn=process_fn, job_repo=job_repo, batch_repo=batch_repo,
        job_label_fn=lambda it: str(it), cleanup_fn=cleanup_fn,
    )


def test_cleanup_fn_runs_once_after_completion(job_repo, batch_repo):
    from app.workers.batch_job_worker import BatchOutcome
    calls = []
    worker = _make_worker(job_repo, batch_repo,
                           lambda items: BatchOutcome(success=True, success_count=len(items)),
                           cleanup_fn=lambda: calls.append(1))
    worker.run()
    assert calls == [1]


def test_cleanup_fn_runs_on_worker_thread_not_caller(job_repo, batch_repo):
    """Playwright sync API 物件綁定建立執行緒：收尾必須在 worker 執行緒上執行"""
    from app.workers.batch_job_worker import BatchOutcome
    thread_ids = {}

    def cleanup():
        thread_ids["cleanup"] = threading.get_ident()

    worker = _make_worker(job_repo, batch_repo,
                           lambda items: BatchOutcome(success=True), cleanup_fn=cleanup)
    t = threading.Thread(target=worker.run)
    t.start()
    t.join(timeout=30)
    assert thread_ids["cleanup"] == t.ident
    assert thread_ids["cleanup"] != threading.get_ident()


def test_cleanup_fn_runs_even_when_batches_fail(job_repo, batch_repo):
    calls = []

    def boom(items):
        raise RuntimeError("批次爆炸")

    worker = _make_worker(job_repo, batch_repo, boom, cleanup_fn=lambda: calls.append(1))
    worker.run()  # 例外由批次層捕捉為 retryable，不外拋
    assert calls == [1]


def test_cleanup_fn_runs_on_cancelled_early_return(job_repo, batch_repo):
    """取消提前 return 的路徑也必須收尾（finally 涵蓋所有 return）"""
    from app.workers.batch_job_worker import BatchOutcome
    calls = []
    worker = _make_worker(job_repo, batch_repo,
                           lambda items: BatchOutcome(success=True),
                           cleanup_fn=lambda: calls.append(1))
    worker._cancel = True
    worker.run()
    assert calls == [1]


def test_cleanup_fn_exception_does_not_crash_worker(job_repo, batch_repo):
    from app.workers.batch_job_worker import BatchOutcome

    def bad_cleanup():
        raise RuntimeError("收尾失敗")

    worker = _make_worker(job_repo, batch_repo,
                           lambda items: BatchOutcome(success=True), cleanup_fn=bad_cleanup)
    worker.run()  # 不應外拋
    assert job_repo.get(worker.job_id).status == "completed"


def test_scraping_worker_closes_browser_via_cleanup_fn(news_repo, job_repo, batch_repo):
    """瀏覽器關閉改走 cleanup_fn：worker.run() 結束後瀏覽器已被關閉"""
    from app.models.news import NewsItem
    from app.services.scraping.body_scraper import FetchOutcome
    from app.workers.scraping_worker import build_scraping_worker

    it = NewsItem(row_id="n1", title="測試", url="http://example.com/a", retained=True)
    news_repo.upsert_one(it)

    class FakeBodyScraper:
        def fetch(self, url):
            return FetchOutcome(status="失敗", detail="未取得可用正文（無法辨識乾淨主文容器）")

    class FakeBrowserScraper:
        closed = False

        def fetch(self, url):
            return FetchOutcome(status="成功", detail="瀏覽器渲染",
                                 body_text="這是瀏覽器渲染取得的完整測試正文，" * 10,
                                 quality_score=0.9, word_count=150)

        def close(self):
            FakeBrowserScraper.closed = True

    worker = build_scraping_worker(
        [it], FakeBodyScraper(), news_repo, job_repo, batch_repo,
        browser_scraper_factory=lambda: FakeBrowserScraper())
    assert worker.cleanup_fn is not None  # 收尾走 cleanup_fn，而非 finished_job signal
    worker.run()
    assert FakeBrowserScraper.closed is True
    assert news_repo.get("n1").body_fetch_status == "成功"


# ---------------------------------------------------------------------------
# 修改 3：Word 匯出版型
# ---------------------------------------------------------------------------
def _export_sample(tmp_path, **topic_overrides):
    from app.models.topic import Topic, Stance
    from app.models.news import NewsItem
    from app.models.settings import WordExportSettings
    from app.exporters.word_exporter import export_daily_report

    fields = dict(
        topic_id="t1", topic_name="測試議題",
        summary_150="一五零字摘要內容。", summary_300="三百字摘要內容。",
        development_progress="發展進度內容。", core_disputes="核心爭點內容。",
        key_actors="1. 甲部長：表示支持 2. 乙議員：提出質疑",
        possible_impact="可能後續影響內容。",
        has_identifiable_stance=True,
    )
    fields.update(topic_overrides)
    topic = Topic(**fields)
    news = [NewsItem(row_id="r1", title="超連結測試新聞", source="中央社",
                      published_at="2026-07-01", url="http://example.com/1", body_text="正文")]
    stance = Stance(stance_id="s1", topic_id="t1", stance_type="支持", speaker="甲部長",
                     organization="某部", claim="支持政策", evidence_news_id="r1",
                     evidence_excerpt="表示支持")
    out = tmp_path / "report.docx"
    export_daily_report(str(out), [topic], {"t1": news}, {"t1": [stance]}, [],
                         WordExportSettings())
    from docx import Document
    return Document(str(out))


def test_word_export_no_label_prefixes_and_all_fields_present(tmp_path):
    doc = _export_sample(tmp_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # 標籤前綴移除，只印內容
    assert "150 字摘要：" not in full_text
    assert "核心爭點：" not in full_text
    assert "一五零字摘要內容。" in full_text
    assert "核心爭點內容。" in full_text
    # 補輸出 key_actors 與 possible_impact
    assert "甲部長" in full_text
    assert "可能後續影響內容。" in full_text


def test_word_export_empty_fields_omitted(tmp_path):
    doc = _export_sample(tmp_path, summary_300="", summary_full="",
                          development_progress="", possible_impact="")
    texts = [p.text for p in doc.paragraphs]
    assert "" not in [t for t in texts if t is None]  # 防呆
    assert not any(t.strip() == "possible_impact" for t in texts)
    full_text = "\n".join(texts)
    assert "一五零字摘要內容。" in full_text


def test_word_export_key_actors_one_per_line_with_numbering_split(tmp_path):
    from app.exporters.word_exporter import _split_actor_lines
    # 單行含編號 → 自動斷行
    assert _split_actor_lines("1. 甲部長：支持 2. 乙議員：質疑") == [
        "1. 甲部長：支持", "2. 乙議員：質疑"]
    # 已是多行 → 維持
    assert _split_actor_lines("甲部長：支持\n乙議員：質疑") == ["甲部長：支持", "乙議員：質疑"]
    assert _split_actor_lines("") == []


def test_word_export_title_hyperlink_no_separate_url_line(tmp_path):
    doc = _export_sample(tmp_path)
    xml = doc.element.xml
    assert "hyperlink" in xml            # 新聞標題為超連結
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "http://example.com/1" not in full_text  # 不另列網址


def test_word_export_cited_news_heading_collapsed(tmp_path):
    doc = _export_sample(tmp_path)
    assert "collapsed" in doc.element.xml  # w15:collapsed


def test_word_export_east_asia_font_applied(tmp_path):
    doc = _export_sample(tmp_path)
    from docx.oxml.ns import qn
    normal_rfonts = doc.styles["Normal"].element.get_or_add_rPr().get_or_add_rFonts()
    assert normal_rfonts.get(qn("w:eastAsia")) == "標楷體"
    heading_rfonts = doc.styles["Heading 1"].element.get_or_add_rPr().get_or_add_rFonts()
    assert heading_rfonts.get(qn("w:eastAsia")) == "標楷體"


def test_word_export_default_font_is_kai():
    from app.models.settings import WordExportSettings
    assert WordExportSettings().font_name == "標楷體"


def test_word_export_actor_section_shown_without_stances(tmp_path):
    """有 key_actors 但無立場條目時，「主要論述與立場」仍應輸出行動者清單"""
    doc = _export_sample(tmp_path, has_identifiable_stance=False)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "主要論述與立場" in full_text
    assert "甲部長" in full_text


# ---------------------------------------------------------------------------
# 修改 6：議題調整頁正文編輯（human_edit_body 資料流）
# ---------------------------------------------------------------------------
def test_human_edit_body_flow_sets_success_and_logs_feedback(news_repo, feedback_repo):
    from app.models.news import NewsItem
    from app.services.feedback.feedback_service import log_feedback
    from app.utils.text_utils import word_count_cjk_aware

    it = NewsItem(row_id="n1", title="抓取失敗新聞", retained=True,
                   body_fetch_status="失敗", body_fetch_detail="404")
    news_repo.upsert_one(it)

    # 模擬頁面「儲存正文修改」的資料流
    new_body = "人工補完的完整正文內容，" * 10
    news_repo.update_fields("n1", {
        "body_text": new_body, "body_source": "人工編輯正文",
        "body_fetch_status": "成功", "body_fetch_detail": "人工編輯/補完正文",
        "body_fetched_at": time.time(),
        "body_word_count": word_count_cjk_aware(new_body), "body_quality_score": 1.0,
    })
    log_feedback(feedback_repo, batch_id="", entity_type="scraping", entity_id="n1",
                  ai_original_value="狀態：失敗", human_final_value="人工編輯正文（狀態改為成功）",
                  action="human_edit_body", operator="user")

    updated = news_repo.get("n1")
    assert updated.body_fetch_status == "成功"
    assert updated.body_text == new_body

    entries = feedback_repo.list_all(entity_type="scraping")
    assert any(e.action == "human_edit_body" and e.entity_id == "n1" for e in entries)
