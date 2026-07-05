"""
Word 早報匯出 — 對應規格書 十四

使用 python-docx。每個議題輸出：
    議題名稱 / 新聞數量與時間範圍 / 150字摘要 / 300字或完整摘要 /
    事件發展與關鍵進度 / 核心爭點 / 可能後續影響 /
    主要論述與立場（行動者清單＋立場條目，僅在有內容時顯示）/
    引用新聞清單（標題為超連結，預設摺疊）

版型要求（V4.2.1，依實際使用回饋定案）：
    - 全文標楷體：中文字型須另設 w:eastAsia，且標題/清單樣式一併套用，
      否則 Word 對 CJK 字元會回退為預設中文字型
    - 內容段落不印「150 字摘要：」等標籤前綴，只印內容；空欄位自動省略
    - 引用新聞清單標題預設摺疊（w15:collapsed）；新聞標題本身為超連結，不另列網址
    - key_actors（主要行動者與發言）每位一行；單行含編號時自動斷行

可設定：Logo、頁首頁尾、日期格式、字型、字級、標題樣式、段落間距、
是否附新聞連結、是否附正文證據摘錄、是否輸出未取得正文新聞清單。
"""
from __future__ import annotations

import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict

from app.models.topic import Topic, Stance
from app.models.news import NewsItem
from app.models.settings import WordExportSettings
from app.utils.logging_setup import get_logger

logger = get_logger("word_exporter")

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"

# 套用中文字型的樣式（標題與清單樣式不套用會回退為 Word 預設中文字型）
_CJK_STYLE_NAMES = ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3",
                     "List Bullet", "List Number")


def export_daily_report(
    output_path: str,
    topics: List[Topic],
    news_by_topic: Dict[str, List[NewsItem]],
    stances_by_topic: Dict[str, List[Stance]],
    missing_body_news: List[NewsItem],
    settings: WordExportSettings,
) -> str:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- 全域字型設定（含 w:eastAsia 中文字型，標題/清單樣式一併套用）----
    doc.styles["Normal"].font.size = Pt(settings.font_size_pt)
    for style_name in _CJK_STYLE_NAMES:
        try:
            _apply_cjk_font(doc.styles[style_name], settings.font_name)
        except KeyError:
            continue  # 部分範本可能缺少某些樣式，跳過即可

    # ---- 頁首：Logo + 標題 + 日期 ----
    if settings.logo_path and Path(settings.logo_path).exists():
        try:
            doc.add_picture(settings.logo_path, width=Cm(3))
        except Exception as e:
            logger.warning(f"Logo 插入失敗: {e}")

    title_text = settings.header_text or "新聞輿情早報"
    date_str = datetime.now().strftime(settings.date_format)
    h = doc.add_heading(f"{title_text}（{date_str}）", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"本期共納入 {len(topics)} 個議題，涵蓋新聞 "
                       f"{sum(len(v) for v in news_by_topic.values())} 則。")

    # ---- 逐議題輸出 ----
    for idx, topic in enumerate(topics, start=1):
        items = news_by_topic.get(topic.topic_id, [])
        stances = stances_by_topic.get(topic.topic_id, [])

        doc.add_heading(f"{idx}. {topic.topic_name}", level=1)

        time_range = _compute_time_range(items)
        meta = doc.add_paragraph()
        meta.add_run(f"新聞數量：{len(items)} 則　時間範圍：{time_range}").italic = True

        # 內容段落：不印標籤前綴，只印內容；空欄位自動省略
        _add_content_paragraph(doc, topic.summary_150, settings)
        _add_content_paragraph(doc, topic.summary_300 or topic.summary_full, settings)
        _add_content_paragraph(doc, topic.development_progress, settings)
        _add_content_paragraph(doc, topic.core_disputes, settings)
        _add_content_paragraph(doc, topic.possible_impact, settings)

        # 主要論述與立場 = 行動者清單（每位一行）+ 立場條目
        key_actors = getattr(topic, "key_actors", "") or ""
        has_stances = bool(topic.has_identifiable_stance and stances)
        if key_actors or has_stances:
            doc.add_heading("主要論述與立場", level=2)
            for line in _split_actor_lines(key_actors):
                doc.add_paragraph(line, style="List Bullet")
            if has_stances:
                for label in ("支持", "反對／質疑", "官方回應"):
                    group = [s for s in stances if s.stance_type == label]
                    if not group:
                        continue
                    doc.add_heading(label, level=3)
                    for s in group:
                        p = doc.add_paragraph(style="List Bullet")
                        who = s.speaker + (f"（{s.organization}）" if s.organization else "")
                        p.add_run(f"{who}：").bold = True
                        p.add_run(s.claim)
                        if settings.include_body_excerpts and s.evidence_excerpt:
                            note = doc.add_paragraph()
                            note.add_run(f"　證據摘錄：{s.evidence_excerpt}").italic = True

        # 引用新聞清單：標題預設摺疊；新聞標題本身為超連結，不另列網址
        cite_heading = doc.add_heading("引用新聞清單", level=2)
        _set_heading_collapsed(cite_heading)
        for it in items:
            p = doc.add_paragraph(style="List Number")
            if settings.include_news_links and it.url:
                _add_hyperlink(p, it.url, it.title, settings.font_name)
            else:
                p.add_run(it.title)
            p.add_run(f"（{it.source}，{it.published_at}）")

    # ---- 未取得正文新聞清單 ----
    if settings.include_missing_body_list and missing_body_news:
        doc.add_heading("附錄：未取得可用正文之新聞清單", level=1)
        for it in missing_body_news:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{it.title}（{it.source}，{it.published_at}） "
                       f"— 狀態：{it.body_fetch_status}；原因：{it.body_fetch_detail}")

    # ---- 頁尾 ----
    if settings.footer_text:
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = settings.footer_text

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info(f"Word 早報已輸出: {out_path}")
    return str(out_path)


def export_simple_topic_list(
    output_path: str,
    topics: List[Topic],
    news_by_topic: Dict[str, List[NewsItem]],
    settings: WordExportSettings,
) -> str:
    """簡易清單匯出：議題名稱當標題分組，底下每則新聞「標題＋連結」各一行。
    不含摘要、立場、統計等內容——供快速分享/剪貼用，與完整早報並存。"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(settings.font_size_pt)
    for style_name in _CJK_STYLE_NAMES:
        try:
            _apply_cjk_font(doc.styles[style_name], settings.font_name)
        except KeyError:
            continue

    idx = 0
    for topic in topics:
        items = news_by_topic.get(topic.topic_id, [])
        if not items:
            continue  # 空議題跳過（不佔編號）
        idx += 1
        doc.add_heading(f"{idx}. {topic.topic_name}", level=1)
        for it in items:
            # 標題行格式：「新聞來源-標題」；來源為空時只印標題
            doc.add_paragraph(f"{it.source}-{it.title}" if it.source else it.title)
            if it.url:
                p = doc.add_paragraph()
                _add_hyperlink(p, it.url, it.url, settings.font_name)

    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info(f"簡易議題清單已輸出: {out_path}")
    return str(out_path)


def _apply_cjk_font(style, font_name: str) -> None:
    """同時設定西文字型與 w:eastAsia 中文字型。
    python-docx 的 style.font.name 只設 w:ascii/w:hAnsi，CJK 字元會回退為
    Word 預設中文字型，必須另外寫入 rFonts 的 w:eastAsia 屬性。"""
    from docx.oxml.ns import qn
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_heading_collapsed(paragraph) -> None:
    """引用新聞清單標題預設摺疊（w15:collapsed）。
    w15 為 Word 2013+ 擴充命名空間，舊版 Word 依 mc:Ignorable 規則忽略，無害。"""
    try:
        from lxml import etree
        ppr = paragraph._p.get_or_add_pPr()
        collapsed = etree.SubElement(ppr, f"{{{_W15_NS}}}collapsed")
        collapsed.set(f"{{{_W_NS}}}val", "true")
    except Exception as e:
        logger.debug(f"設定標題摺疊失敗（不影響輸出）: {e}")


def _add_hyperlink(paragraph, url: str, text: str, font_name: str) -> None:
    """在段落內加入超連結 run（藍色＋底線，沿用段落中文字型）。
    失敗時退回純文字，確保匯出不因單一網址異常而中斷。"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.opc.constants import RELATIONSHIP_TYPE

        r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), font_name)
        rpr.append(rfonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rpr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rpr.append(underline)
        run.append(rpr)

        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
    except Exception as e:
        logger.warning(f"超連結建立失敗，改為純文字（{url}）: {e}")
        paragraph.add_run(text)


def _add_content_paragraph(doc, text: str, settings: WordExportSettings) -> None:
    """內容段落：只印內容（不加標籤前綴），空欄位自動省略"""
    if not text:
        return
    from docx.shared import Pt
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(settings.paragraph_spacing_pt)


_ACTOR_NUMBERING_RE = re.compile(r"\s*(?=(?:\d{1,2}|[一二三四五六七八九十])[\.、．)）]\s*\S)")


def _split_actor_lines(key_actors: str) -> List[str]:
    """行動者清單每位一行：以換行切分；模型把多位擠在同一行時依編號自動斷行"""
    if not key_actors:
        return []
    lines = [ln.strip() for ln in key_actors.splitlines() if ln.strip()]
    out: List[str] = []
    for line in lines:
        parts = [p.strip() for p in _ACTOR_NUMBERING_RE.split(line) if p.strip()]
        out.extend(parts if len(parts) > 1 else [line])
    return out


def _compute_time_range(items: List[NewsItem]) -> str:
    dates = [it.published_at for it in items if it.published_at]
    if not dates:
        return "無時間資訊"
    dates_sorted = sorted(dates)
    if dates_sorted[0] == dates_sorted[-1]:
        return dates_sorted[0]
    return f"{dates_sorted[0]} ～ {dates_sorted[-1]}"
